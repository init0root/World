
# coding: utf-8

from tkinter import *
from PIL import ImageTk, Image
import os
from docx import Document
from docx.shared import Inches
from tkinter import messagebox
import random
import wget

path = 'sources/la.png'

#configurações da primeira janela
window = Tk()
window.title("world")
window.geometry('350x200')
img = ImageTk.PhotoImage(Image.open(path))
lbl = Label(window, image = img).grid(row=5)
#função que gera a prova
def click():
    
	arquivo = open("worlist.txt", "r")
	ler = arquivo.read()
	documento = Document()
	documento.add_picture('sources/cab.png', width=Inches(6))
	documento.add_paragraph(ler)
	save = random.randint(0,100)								
	documento.save('{}.docx'.format(save))
	

	messagebox.showinfo('world', 'Avaliação gerada com sucesso')

#Configurações da 2° Janela

#função que baixa as provas do primeiro ano
def click_in_chk1():
	os.remove('worlist.txt')
	wget.download('http://download1585.mediafire.com/d8yc2o8th0yg/t4exwu5d7acvhhy/worlist.txt')

#função que baixa as provas do segundo ano
def click_in_chk2():
	os.remove('worlist.txt')
	wget.download('http://download1349.mediafire.com/ggf64gannxgg/rkfvgpiwgr9ff34/worlist.txt')

#função que baica as provas do terceiro ano

def click_in_chk3():
	os.remove('worlist.txt')
	wget.download('http://download1503.mediafire.com/bmk1p6xutjtg/fl5vap39s3uooqg/worlist.txt')

def click_in_about():
	window3 = Tk()
	window3.title("Sobre nos")
	window3.geometry("200x200")
	text0 = Label(window3, text="Desenvolvedor: Henrique Moura ", fg='red').grid(row=1)
	text1 = Label(window3, text="Wordlist: Letícia de Jesus Teles", fg='green').grid(row=2)

#função que abre o modo editor
def click2():

	window2 = Tk()
	window2.title("Editor")
	window2.geometry("450x300")
	text = Label(window2, text="Selecione a wordlist para baixar novas questões: ").grid(row=0)
	text2 = Label(window2, text="Para baixar os arquivos é necessario internet.", fg="red").grid(row=4)
	text3 = Label(window2, text="Por enquanto só estão disponiveis questões de Física do 1° ano", fg="red").grid(row=5)
	chk = Checkbutton(window2, text='1° Wordlist', command=click_in_chk1).grid(row=1)
	chk2 = Checkbutton(window2, text='2° Wordlist', command=click_in_chk2).grid(row=2)
	chk3 = Checkbutton(window2, text='3° Wordlist', command=click_in_chk3).grid(row=3)


#
	

#botões do primeiro menu, gerar prova e editor
b = Button(window, text="Gerar Prova", fg="green", command=click).grid(row=1)
c = Button(window, text ="Editor", fg="red", command=click2).grid(row=2)
d = Button(window, text="Sobre nós", fg="purple", command=click_in_about).grid(row=3)
e = Button(window, text="Apoie o projeto", fg="blue").grid(row=4)



window.mainloop()